"""
End-to-End Test Script cho Phase 2 Pipeline
=============================================

Script này test toàn bộ pipeline KHÔNG cần MongoDB:
1. Text Extraction (PDF, DOCX, TXT) — tạo file test → extract → verify
2. Text Preprocessing — noise removal, truncation
3. Gemini API — gọi API thật → sinh quiz → validate response

Usage:
    cd backend
    python -m tests.test_e2e_pipeline

Yêu cầu:
    - GEMINI_API_KEY hợp lệ trong .env
    - KHÔNG cần Docker/MongoDB (test này bỏ qua DB layer)
"""

import asyncio
import sys
import os
from io import BytesIO
from datetime import datetime

# Fix Windows console encoding
if sys.platform == "win32":
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    sys.stderr.reconfigure(encoding="utf-8", errors="replace")

# Add project root to path
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from app.services.text_extractor import TextExtractor
from app.services.ai_service import DeepSeekService, DeepSeekServiceError
from app.utils.text_processing import preprocess_text, estimate_tokens
from app.config import get_settings


# ============================================
# ANSI Colors for pretty output
# ============================================
class Colors:
    GREEN = "\033[92m"
    RED = "\033[91m"
    YELLOW = "\033[93m"
    BLUE = "\033[94m"
    CYAN = "\033[96m"
    BOLD = "\033[1m"
    END = "\033[0m"


def ok(msg): print(f"  {Colors.GREEN}✓{Colors.END} {msg}")
def fail(msg): print(f"  {Colors.RED}✗{Colors.END} {msg}")
def info(msg): print(f"  {Colors.BLUE}ℹ{Colors.END} {msg}")
def header(msg): print(f"\n{Colors.BOLD}{Colors.CYAN}{'='*60}\n  {msg}\n{'='*60}{Colors.END}")
def subheader(msg): print(f"\n  {Colors.BOLD}{Colors.YELLOW}▸ {msg}{Colors.END}")


# ============================================
# TEST 1: Text Extraction
# ============================================
def test_text_extraction():
    """Test text extraction cho cả 3 file types."""
    header("TEST 1: Text Extraction Service")
    extractor = TextExtractor()
    results = {}

    # --- TXT ---
    subheader("TXT Extraction")
    txt_content = (
        "Python là ngôn ngữ lập trình bậc cao, đa mục đích.\n"
        "Python được thiết kế bởi Guido van Rossum và ra mắt năm 1991.\n"
        "Python hỗ trợ lập trình hướng đối tượng, lập trình hàm.\n"
        "Các thư viện phổ biến: NumPy, Pandas, TensorFlow, FastAPI.\n"
        "Python có cú pháp đơn giản, dễ đọc, dễ học.\n"
    )
    txt_bytes = txt_content.encode("utf-8")
    text = extractor.extract(txt_bytes, "txt")
    is_valid, msg = extractor.validate_extracted_text(text)

    if is_valid and "Python" in text and "Guido van Rossum" in text:
        ok(f"TXT extracted successfully: {len(text)} chars")
        results["txt"] = text
    else:
        fail(f"TXT extraction failed: {msg}")
        return None

    # --- DOCX ---
    subheader("DOCX Extraction")
    try:
        from docx import Document
        doc = Document()
        doc.add_heading("Machine Learning cơ bản", level=1)
        doc.add_paragraph(
            "Machine Learning (ML) là một nhánh của Trí tuệ nhân tạo (AI), "
            "cho phép máy tính học từ dữ liệu mà không cần lập trình tường minh."
        )
        doc.add_paragraph(
            "Có 3 loại ML chính: Supervised Learning (học có giám sát), "
            "Unsupervised Learning (học không giám sát), và Reinforcement Learning (học tăng cường)."
        )
        doc.add_paragraph(
            "Supervised Learning gồm các thuật toán như Linear Regression, "
            "Decision Tree, Random Forest, Support Vector Machine, và Neural Networks."
        )
        doc.add_paragraph(
            "Deep Learning là nhánh con của Machine Learning, sử dụng neural networks "
            "nhiều lớp để học các biểu diễn phức tạp từ dữ liệu."
        )

        buf = BytesIO()
        doc.save(buf)
        docx_bytes = buf.getvalue()

        text = extractor.extract(docx_bytes, "docx")
        is_valid, msg = extractor.validate_extracted_text(text)

        if is_valid and "Machine Learning" in text:
            ok(f"DOCX extracted successfully: {len(text)} chars")
            results["docx"] = text
        else:
            fail(f"DOCX extraction failed: {msg}")
    except Exception as e:
        fail(f"DOCX extraction error: {e}")

    # --- PDF ---
    subheader("PDF Extraction")
    try:
        import fitz
        pdf_doc = fitz.open()
        page = pdf_doc.new_page()
        pdf_text = (
            "Giới thiệu về FastAPI\n\n"
            "FastAPI là một web framework hiện đại, hiệu suất cao cho Python, "
            "được xây dựng trên Starlette và Pydantic.\n\n"
            "FastAPI tự động tạo API documentation (Swagger UI và ReDoc) "
            "từ type hints trong code Python.\n\n"
            "FastAPI hỗ trợ async/await, giúp xử lý nhiều request đồng thời "
            "với hiệu suất tương đương Node.js và Go.\n\n"
            "Pydantic là thư viện validation dữ liệu, tự động kiểm tra "
            "kiểu dữ liệu và chuyển đổi format."
        )
        rect = fitz.Rect(72, 72, 540, 720)
        page.insert_textbox(rect, pdf_text, fontsize=11)

        buf = BytesIO()
        pdf_doc.save(buf)
        pdf_doc.close()
        pdf_bytes = buf.getvalue()

        text = extractor.extract(pdf_bytes, "pdf")
        is_valid, msg = extractor.validate_extracted_text(text)

        if is_valid and "FastAPI" in text:
            ok(f"PDF extracted successfully: {len(text)} chars")
            results["pdf"] = text
        else:
            fail(f"PDF extraction failed: {msg}")
    except Exception as e:
        fail(f"PDF extraction error: {e}")

    return results


# ============================================
# TEST 2: Text Preprocessing
# ============================================
def test_preprocessing(extracted_texts: dict):
    """Test text preprocessing pipeline."""
    header("TEST 2: Text Preprocessing & Chunking")

    for file_type, text in extracted_texts.items():
        subheader(f"Preprocessing {file_type.upper()} text")

        processed = preprocess_text(text)
        tokens = estimate_tokens(processed)

        info(f"Original:    {len(text)} chars")
        info(f"Processed:   {len(processed)} chars")
        info(f"Est. tokens: {tokens}")

        if len(processed) > 0:
            ok(f"{file_type.upper()} preprocessing successful")
        else:
            fail(f"{file_type.upper()} preprocessing returned empty text")

    # Combine all texts for Gemini test
    combined = "\n\n".join(extracted_texts.values())
    processed = preprocess_text(combined)
    info(f"\nCombined text: {len(combined)} chars → {len(processed)} chars processed")

    return processed


# ============================================
# TEST 3: DeepSeek AI Quiz Generation
# ============================================
async def test_deepseek_generation(text: str):
    """Test DeepSeek AI quiz generation pipeline."""
    header("TEST 3: DeepSeek AI Quiz Generation")

    settings = get_settings()

    # Initialize service
    subheader("Initializing DeepSeekService")
    try:
        service = DeepSeekService(api_key=settings.DEEPSEEK_API_KEY)
        ok("DeepSeekService initialized successfully")
    except ValueError as e:
        fail(f"DeepSeekService init failed: {e}")
        return None

    # Generate quiz
    subheader("Generating quiz (5 questions, difficulty=mixed, lang=vi)")
    info("Calling DeepSeek API... (this may take 15-30 seconds)")

    start_time = datetime.now()
    try:
        quiz_data = await service.generate_quiz(
            text=text,
            num_questions=5,
            difficulty="mixed",
            language="vi",
        )
        elapsed = (datetime.now() - start_time).total_seconds()
        ok(f"Quiz generated in {elapsed:.1f}s")
    except DeepSeekServiceError as e:
        elapsed = (datetime.now() - start_time).total_seconds()
        fail(f"Quiz generation failed after {elapsed:.1f}s: {e}")
        return None
    except Exception as e:
        fail(f"Unexpected error: {e}")
        return None

    # Validate response structure
    subheader("Validating quiz structure")

    # Check title
    title = quiz_data.get("title", "")
    if title:
        ok(f"Title: \"{title}\"")
    else:
        fail("Missing title")

    # Check description
    desc = quiz_data.get("description", "")
    if desc:
        ok(f"Description: \"{desc[:80]}...\"")
    else:
        info("No description (optional)")

    # Check questions
    questions = quiz_data.get("questions", [])
    ok(f"Questions generated: {len(questions)}")

    if not questions:
        fail("No questions generated!")
        return quiz_data

    # Validate each question
    subheader(f"Validating {len(questions)} questions")
    all_valid = True

    for i, q in enumerate(questions):
        q_text = q.get("question_text", "")[:60]
        options = q.get("options", [])
        correct = q.get("correct_answer", "")
        explanation = q.get("explanation", "")

        errors = []
        if not q.get("question_text"):
            errors.append("missing question_text")
        if len(options) != 4:
            errors.append(f"has {len(options)} options (expected 4)")
        if correct not in {"A", "B", "C", "D"}:
            errors.append(f"invalid correct_answer: '{correct}'")
        
        option_keys = {o.get("key") for o in options}
        if option_keys != {"A", "B", "C", "D"}:
            errors.append(f"invalid option keys: {option_keys}")

        empty_opts = [o for o in options if not o.get("text", "").strip()]
        if empty_opts:
            errors.append(f"{len(empty_opts)} empty option text(s)")

        if errors:
            fail(f"Q{i+1}: {', '.join(errors)}")
            all_valid = False
        else:
            ok(f"Q{i+1}: \"{q_text}...\" [Answer: {correct}]")

    # Print a sample question in detail
    subheader("Sample Question (Detail)")
    sample = questions[0]
    print(f"\n  {Colors.BOLD}📝 {sample['question_text']}{Colors.END}")
    for opt in sample["options"]:
        marker = "→" if opt["key"] == sample["correct_answer"] else " "
        color = Colors.GREEN if opt["key"] == sample["correct_answer"] else ""
        end = Colors.END if color else ""
        print(f"    {color}{marker} {opt['key']}. {opt['text']}{end}")
    if sample.get("explanation"):
        print(f"\n  {Colors.CYAN}💡 {sample['explanation']}{Colors.END}")

    return quiz_data


# ============================================
# MAIN
# ============================================
async def main():
    print(f"\n{Colors.BOLD}{Colors.CYAN}")
    print("╔══════════════════════════════════════════════════════════╗")
    print("║     AIQuiz — Phase 2 End-to-End Pipeline Test           ║")
    print("╚══════════════════════════════════════════════════════════╝")
    print(f"{Colors.END}")

    total_tests = 0
    passed_tests = 0

    # TEST 1: Text Extraction
    total_tests += 1
    extracted = test_text_extraction()
    if extracted:
        passed_tests += 1

    # TEST 2: Preprocessing
    if extracted:
        total_tests += 1
        processed_text = test_preprocessing(extracted)
        if processed_text:
            passed_tests += 1

        # TEST 3: DeepSeek AI
        total_tests += 1
        quiz = await test_deepseek_generation(processed_text)
        if quiz and quiz.get("questions"):
            passed_tests += 1

    # Summary
    header("SUMMARY")
    color = Colors.GREEN if passed_tests == total_tests else Colors.RED
    print(f"\n  {color}{Colors.BOLD}Result: {passed_tests}/{total_tests} tests passed{Colors.END}")

    if passed_tests == total_tests:
        print(f"\n  {Colors.GREEN}{Colors.BOLD}🎉 Phase 2 Pipeline: ALL TESTS PASSED!{Colors.END}")
        print(f"  {Colors.GREEN}Pipeline sẵn sàng: Upload → Extract → Preprocess → AI → Quiz{Colors.END}\n")
    else:
        print(f"\n  {Colors.RED}⚠ Some tests failed. Check errors above.{Colors.END}\n")


if __name__ == "__main__":
    asyncio.run(main())
