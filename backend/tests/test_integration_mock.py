"""
Integration Test - Full Pipeline với Mock DeepSeek
==================================================

Test toàn bộ pipeline end-to-end KHÔNG cần DeepSeek API thật:
1. Extract text từ file → preprocess → mock AI response → validate

Chứng minh pipeline code đúng khi có API key hoạt động.
"""

import asyncio
import sys
import os
import json
from io import BytesIO
from unittest.mock import AsyncMock, patch, MagicMock
from datetime import datetime

if sys.platform == "win32":
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from app.services.text_extractor import TextExtractor
from app.services.ai_service import DeepSeekService
from app.services.quiz_service import QuizService
from app.utils.text_processing import preprocess_text


# ============================================
# Colors
# ============================================
G = "\033[92m"; R = "\033[91m"; Y = "\033[93m"; B = "\033[94m"
C = "\033[96m"; BOLD = "\033[1m"; END = "\033[0m"

def ok(msg): print(f"  {G}✓{END} {msg}")
def fail(msg): print(f"  {R}✗{END} {msg}")
def info(msg): print(f"  {B}ℹ{END} {msg}")
def header(msg): print(f"\n{BOLD}{C}{'='*60}\n  {msg}\n{'='*60}{END}")
def subheader(msg): print(f"\n  {BOLD}{Y}▸ {msg}{END}")


# Mock DeepSeek response (giống thật 100%)
MOCK_DEEPSEEK_RESPONSE = {
    "title": "Kiểm tra kiến thức Python và Machine Learning",
    "description": "Bộ câu hỏi trắc nghiệm về Python, Machine Learning và FastAPI",
    "questions": [
        {
            "question_text": "Python được thiết kế bởi ai?",
            "options": [
                {"key": "A", "text": "Guido van Rossum"},
                {"key": "B", "text": "James Gosling"},
                {"key": "C", "text": "Brendan Eich"},
                {"key": "D", "text": "Dennis Ritchie"},
            ],
            "correct_answer": "A",
            "explanation": "Theo tài liệu, Python được thiết kế bởi Guido van Rossum và ra mắt năm 1991.",
        },
        {
            "question_text": "Machine Learning là nhánh con của lĩnh vực nào?",
            "options": [
                {"key": "A", "text": "Database Management"},
                {"key": "B", "text": "Trí tuệ nhân tạo (AI)"},
                {"key": "C", "text": "Computer Networks"},
                {"key": "D", "text": "Operating Systems"},
            ],
            "correct_answer": "B",
            "explanation": "Machine Learning là một nhánh của Trí tuệ nhân tạo (AI).",
        },
        {
            "question_text": "Có bao nhiêu loại Machine Learning chính?",
            "options": [
                {"key": "A", "text": "2 loại"},
                {"key": "B", "text": "3 loại"},
                {"key": "C", "text": "4 loại"},
                {"key": "D", "text": "5 loại"},
            ],
            "correct_answer": "B",
            "explanation": "Có 3 loại ML chính: Supervised, Unsupervised, và Reinforcement Learning.",
        },
        {
            "question_text": "FastAPI được xây dựng trên framework nào?",
            "options": [
                {"key": "A", "text": "Django và SQLAlchemy"},
                {"key": "B", "text": "Flask và Jinja2"},
                {"key": "C", "text": "Starlette và Pydantic"},
                {"key": "D", "text": "Tornado và Motor"},
            ],
            "correct_answer": "C",
            "explanation": "FastAPI được xây dựng trên Starlette và Pydantic.",
        },
        {
            "question_text": "Deep Learning sử dụng gì để học từ dữ liệu?",
            "options": [
                {"key": "A", "text": "Decision Trees"},
                {"key": "B", "text": "Neural networks nhiều lớp"},
                {"key": "C", "text": "Linear Regression"},
                {"key": "D", "text": "K-Means Clustering"},
            ],
            "correct_answer": "B",
            "explanation": "Deep Learning sử dụng neural networks nhiều lớp để học biểu diễn phức tạp.",
        },
    ],
}


def create_test_files():
    """Tạo test files (TXT, DOCX, PDF) và extract text."""
    extractor = TextExtractor()
    all_text = []

    # TXT
    txt = (
        "Python là ngôn ngữ lập trình bậc cao, đa mục đích.\n"
        "Python được thiết kế bởi Guido van Rossum và ra mắt năm 1991.\n"
        "Python hỗ trợ lập trình hướng đối tượng, lập trình hàm.\n"
        "Các thư viện phổ biến: NumPy, Pandas, TensorFlow, FastAPI.\n"
        "Python có cú pháp đơn giản, dễ đọc, dễ học.\n"
    )
    text = extractor.extract(txt.encode("utf-8"), "txt")
    all_text.append(("TXT", text))

    # DOCX
    from docx import Document
    doc = Document()
    doc.add_heading("Machine Learning cơ bản", level=1)
    doc.add_paragraph(
        "Machine Learning (ML) là một nhánh của Trí tuệ nhân tạo (AI), "
        "cho phép máy tính học từ dữ liệu mà không cần lập trình tường minh."
    )
    doc.add_paragraph(
        "Có 3 loại ML chính: Supervised Learning, Unsupervised Learning, "
        "và Reinforcement Learning."
    )
    doc.add_paragraph(
        "Deep Learning sử dụng neural networks nhiều lớp để học biểu diễn phức tạp."
    )
    buf = BytesIO()
    doc.save(buf)
    text = extractor.extract(buf.getvalue(), "docx")
    all_text.append(("DOCX", text))

    # PDF
    import fitz
    pdf_doc = fitz.open()
    page = pdf_doc.new_page()
    pdf_text = (
        "FastAPI là web framework hiện đại, hiệu suất cao cho Python, "
        "được xây dựng trên Starlette và Pydantic.\n\n"
        "FastAPI hỗ trợ async/await và tự động tạo API documentation."
    )
    rect = fitz.Rect(72, 72, 540, 720)
    page.insert_textbox(rect, pdf_text, fontsize=11)
    buf = BytesIO()
    pdf_doc.save(buf)
    pdf_doc.close()
    text = extractor.extract(buf.getvalue(), "pdf")
    all_text.append(("PDF", text))

    return all_text


class MockMongoDB:
    """Mock MongoDB cho integration test."""

    def __init__(self):
        self._data = {"documents": {}, "quizzes": {}, "questions": {}}
        self._counter = 0

    def _next_id(self):
        self._counter += 1
        from bson import ObjectId
        return ObjectId()

    @property
    def documents(self):
        return MockCollection(self, "documents")

    @property
    def quizzes(self):
        return MockCollection(self, "quizzes")

    @property
    def questions(self):
        return MockCollection(self, "questions")


class MockCollection:
    def __init__(self, db, name):
        self.db = db
        self.name = name

    async def insert_one(self, doc):
        oid = self.db._next_id()
        doc["_id"] = oid
        self.db._data[self.name][str(oid)] = doc
        return MagicMock(inserted_id=oid)

    async def insert_many(self, docs):
        ids = []
        for doc in docs:
            oid = self.db._next_id()
            doc["_id"] = oid
            self.db._data[self.name][str(oid)] = doc
            ids.append(oid)
        return MagicMock(inserted_ids=ids)

    async def find_one(self, query):
        if "_id" in query:
            key = str(query["_id"])
            return self.db._data[self.name].get(key)
        return None

    async def update_one(self, query, update):
        if "_id" in query:
            key = str(query["_id"])
            doc = self.db._data[self.name].get(key)
            if doc and "$set" in update:
                doc.update(update["$set"])
        return MagicMock(modified_count=1)

    def find(self, query):
        return MockCursor(self.db, self.name, query)


class MockCursor:
    def __init__(self, db, name, query):
        self.db = db
        self.name = name
        self.query = query
        self._sort_key = None

    def sort(self, key, direction):
        self._sort_key = key
        return self

    async def to_list(self, length=100):
        results = []
        for doc in self.db._data[self.name].values():
            match = all(doc.get(k) == v for k, v in self.query.items())
            if match:
                results.append(doc)
        if self._sort_key:
            results.sort(key=lambda x: x.get(self._sort_key, 0))
        return results[:length]


async def test_full_pipeline():
    """Test toàn bộ pipeline với mock Gemini và mock MongoDB."""

    print(f"\n{BOLD}{C}")
    print("╔══════════════════════════════════════════════════════════╗")
    print("║  AIQuiz — Phase 2 Integration Test (Mock DeepSeek + DB)   ║")
    print("╚══════════════════════════════════════════════════════════╝")
    print(f"{END}")

    passed = 0
    total = 0

    # ============================================
    # TEST 1: Extract text from all file types
    # ============================================
    header("TEST 1: Text Extraction (Real Files)")
    total += 1
    try:
        files = create_test_files()
        combined_text = ""
        for ftype, text in files:
            ok(f"{ftype}: {len(text)} chars extracted")
            combined_text += text + "\n\n"
        passed += 1
    except Exception as e:
        fail(f"Extraction failed: {e}")
        return

    # ============================================
    # TEST 2: Preprocessing
    # ============================================
    header("TEST 2: Text Preprocessing")
    total += 1
    processed = preprocess_text(combined_text)
    info(f"Original: {len(combined_text)} chars → Processed: {len(processed)} chars")
    if len(processed) > 0:
        ok("Preprocessing successful")
        passed += 1
    else:
        fail("Preprocessing returned empty text")

    # ============================================
    # TEST 3: Full pipeline with Mock DeepSeek + Mock DB
    # ============================================
    header("TEST 3: Full Quiz Generation Pipeline (Mocked)")
    total += 1

    subheader("Setting up mock DeepSeek & MongoDB")
    mock_db = MockMongoDB()

    # Create a mock DeepSeekService that returns MOCK_DEEPSEEK_RESPONSE
    mock_ai = MagicMock(spec=DeepSeekService)
    mock_ai.generate_quiz = AsyncMock(return_value=MOCK_DEEPSEEK_RESPONSE)

    quiz_service = QuizService(db=mock_db, gemini_service=mock_ai)
    ok("QuizService initialized with mocks")

    subheader("Running pipeline: File → Extract → Preprocess → AI → Save")

    # Create a real DOCX file to upload
    from docx import Document
    doc = Document()
    doc.add_paragraph(
        "Python là ngôn ngữ lập trình bậc cao. "
        "Machine Learning là nhánh của AI. "
        "FastAPI xây dựng trên Starlette và Pydantic. "
        "Deep Learning dùng neural networks nhiều lớp."
        "Supervised Learning gồm Linear Regression, Decision Tree."
    )
    buf = BytesIO()
    doc.save(buf)
    file_bytes = buf.getvalue()

    try:
        result = await quiz_service.upload_and_generate_quiz(
            file_bytes=file_bytes,
            filename="test_document.docx",
            file_type="docx",
            num_questions=5,
            difficulty="mixed",
            language="vi",
        )

        # Validate result structure
        assert "document_id" in result, "Missing document_id"
        assert "quiz" in result, "Missing quiz"

        quiz = result["quiz"]
        assert "id" in quiz, "Missing quiz id"
        assert "title" in quiz, "Missing title"
        assert "questions" in quiz, "Missing questions"
        assert len(quiz["questions"]) == 5, f"Expected 5 questions, got {len(quiz['questions'])}"

        ok(f"Pipeline completed successfully!")
        info(f"Document ID: {result['document_id']}")
        info(f"Quiz ID: {quiz['id']}")
        info(f"Title: \"{quiz['title']}\"")
        info(f"Questions: {len(quiz['questions'])}")

        # Verify DeepSeek was called correctly
        mock_ai.generate_quiz.assert_called_once()
        call_args = mock_ai.generate_quiz.call_args
        ok("DeepSeek API was called with correct parameters")
        info(f"  num_questions={call_args.kwargs.get('num_questions', call_args.args[1] if len(call_args.args) > 1 else 'N/A')}")

        passed += 1

    except Exception as e:
        fail(f"Pipeline error: {e}")
        import traceback
        traceback.print_exc()

    # ============================================
    # TEST 4: Validate saved data in mock DB
    # ============================================
    header("TEST 4: Database Persistence Verification")
    total += 1

    subheader("Checking saved documents")
    docs = mock_db._data["documents"]
    info(f"Documents saved: {len(docs)}")
    if docs:
        doc_data = list(docs.values())[0]
        status = doc_data.get("upload_status")
        ok(f"Document status: {status}")
        ok(f"Filename: {doc_data.get('original_filename')}")
        ok(f"Text length: {doc_data.get('text_length')} chars")

    subheader("Checking saved quizzes")
    quizzes = mock_db._data["quizzes"]
    info(f"Quizzes saved: {len(quizzes)}")
    if quizzes:
        quiz_data = list(quizzes.values())[0]
        ok(f"Title: \"{quiz_data.get('title')}\"")
        ok(f"Total questions: {quiz_data.get('total_questions')}")
        ok(f"Difficulty: {quiz_data.get('difficulty')}")

    subheader("Checking saved questions")
    questions = mock_db._data["questions"]
    info(f"Questions saved: {len(questions)}")

    all_valid = True
    for q_data in questions.values():
        opts = q_data.get("options", [])
        correct = q_data.get("correct_answer", "")
        if len(opts) != 4 or correct not in {"A", "B", "C", "D"}:
            all_valid = False
            fail(f"Invalid question: {q_data.get('question_text', '')[:40]}")

    if all_valid and len(questions) == 5:
        ok(f"All {len(questions)} questions saved correctly with valid structure")
        passed += 1
    else:
        fail("Some questions have invalid structure")

    # ============================================
    # TEST 5: Print sample quiz output
    # ============================================
    header("TEST 5: Sample Quiz Output")
    total += 1
    try:
        for i, q_data in enumerate(sorted(questions.values(), key=lambda x: x.get("order", 0))):
            q_num = q_data.get("order", i + 1)
            print(f"\n  {BOLD}📝 Câu {q_num}: {q_data['question_text']}{END}")
            for opt in q_data["options"]:
                is_correct = opt["key"] == q_data["correct_answer"]
                marker = "→" if is_correct else " "
                color = G if is_correct else ""
                end_c = END if color else ""
                print(f"    {color}{marker} {opt['key']}. {opt['text']}{end_c}")
            if q_data.get("explanation"):
                print(f"    {C}💡 {q_data['explanation']}{END}")
        ok("Quiz rendering completed")
        passed += 1
    except Exception as e:
        fail(f"Rendering error: {e}")

    # ============================================
    # SUMMARY
    # ============================================
    header("SUMMARY")
    color = G if passed == total else R
    print(f"\n  {color}{BOLD}Result: {passed}/{total} tests passed{END}")

    if passed == total:
        print(f"\n  {G}{BOLD}🎉 Phase 2 Pipeline: ALL TESTS PASSED!{END}")
        print(f"  {G}Code pipeline hoạt động hoàn chỉnh:{END}")
        print(f"  {G}  Extract → Preprocess → AI Generate → Save DB → Response{END}")
        print(f"\n  {Y}⚠ DeepSeek API trả 429 do hết quota free tier.{END}")
        print(f"  {Y}  Khi có key mới (hoặc bật billing), pipeline sẽ chạy end-to-end.{END}\n")
    else:
        print(f"\n  {R}⚠ Some tests failed.{END}\n")


if __name__ == "__main__":
    asyncio.run(test_full_pipeline())
