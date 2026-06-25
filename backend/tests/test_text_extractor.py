"""
Test Text Extractor - Unit tests cho text extraction service
==============================================================

Test strategy:
- Tạo file test nhỏ (PDF, DOCX, TXT) programmatically
- Verify extract trả về text đúng
- Test edge cases: file rỗng, encoding lạ, PDF corrupt
"""

import pytest
from io import BytesIO

from app.services.text_extractor import TextExtractor


# ============================================
# FIXTURES
# ============================================

@pytest.fixture
def extractor():
    """Tạo TextExtractor instance."""
    return TextExtractor()


@pytest.fixture
def sample_txt_bytes():
    """Tạo sample TXT file bytes."""
    text = (
        "Đây là nội dung tài liệu mẫu bằng tiếng Việt.\n"
        "Python là ngôn ngữ lập trình phổ biến nhất thế giới.\n"
        "Machine Learning là nhánh con của Artificial Intelligence.\n"
        "FastAPI là web framework hiện đại cho Python.\n"
        "MongoDB là cơ sở dữ liệu NoSQL document-oriented.\n"
    )
    return text.encode("utf-8")


@pytest.fixture
def sample_txt_short():
    """Tạo TXT file quá ngắn (< 100 chars)."""
    return "Hello World".encode("utf-8")


@pytest.fixture
def sample_docx_bytes():
    """Tạo sample DOCX file bytes."""
    from docx import Document

    doc = Document()
    doc.add_heading("Bài giảng Python cơ bản", level=1)
    doc.add_paragraph(
        "Python là ngôn ngữ lập trình bậc cao, được thiết kế bởi "
        "Guido van Rossum và ra mắt lần đầu vào năm 1991."
    )
    doc.add_paragraph(
        "Python hỗ trợ nhiều paradigm: lập trình hướng đối tượng, "
        "lập trình hàm, và lập trình thủ tục."
    )
    doc.add_paragraph(
        "Các thư viện phổ biến: NumPy, Pandas, TensorFlow, FastAPI, Django."
    )

    # Thêm table
    table = doc.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "Library"
    table.cell(0, 1).text = "Use Case"
    table.cell(1, 0).text = "NumPy"
    table.cell(1, 1).text = "Scientific Computing"
    table.cell(2, 0).text = "FastAPI"
    table.cell(2, 1).text = "Web APIs"

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


@pytest.fixture
def sample_pdf_bytes():
    """Tạo sample PDF file bytes."""
    import fitz

    doc = fitz.open()
    page = doc.new_page()

    # Add text to PDF
    text = (
        "Machine Learning Fundamentals\n\n"
        "Machine Learning is a subset of Artificial Intelligence that "
        "enables computers to learn from data without being explicitly "
        "programmed.\n\n"
        "Supervised Learning includes algorithms like Linear Regression, "
        "Decision Trees, and Neural Networks.\n\n"
        "Unsupervised Learning includes Clustering, Dimensionality "
        "Reduction, and Association Rules."
    )
    # Insert text at position
    rect = fitz.Rect(72, 72, 540, 720)
    page.insert_textbox(rect, text, fontsize=12)

    buffer = BytesIO()
    doc.save(buffer)
    doc.close()
    return buffer.getvalue()


# ============================================
# TXT EXTRACTION TESTS
# ============================================

class TestTxtExtraction:
    """Tests cho TXT file extraction."""

    def test_extract_utf8(self, extractor, sample_txt_bytes):
        """Test extract TXT file UTF-8 chuẩn."""
        result = extractor.extract(sample_txt_bytes, "txt")
        assert "Python" in result
        assert "tiếng Việt" in result
        assert len(result) > 0

    def test_extract_txt_with_bom(self, extractor):
        """Test extract TXT file có UTF-8 BOM."""
        bom = b"\xef\xbb\xbf"
        content = bom + "Hello World with BOM marker and enough content to make this longer than minimum".encode("utf-8")
        result = extractor.extract(content, "txt")
        assert "Hello World" in result

    def test_extract_txt_latin1(self, extractor):
        """Test extract TXT file encoding Latin-1."""
        text = "Café résumé naïve"
        content = text.encode("latin-1")
        result = extractor.extract(content, "txt")
        # chardet should detect encoding and decode properly
        assert len(result) > 0

    def test_extract_empty_txt(self, extractor):
        """Test extract TXT file rỗng."""
        result = extractor.extract(b"", "txt")
        assert result == ""


# ============================================
# DOCX EXTRACTION TESTS
# ============================================

class TestDocxExtraction:
    """Tests cho DOCX file extraction."""

    def test_extract_docx(self, extractor, sample_docx_bytes):
        """Test extract DOCX file với paragraphs và tables."""
        result = extractor.extract(sample_docx_bytes, "docx")
        # Check paragraphs content
        assert "Python" in result
        assert "Guido van Rossum" in result
        # Check table content
        assert "NumPy" in result
        assert "FastAPI" in result

    def test_extract_docx_preserves_structure(self, extractor, sample_docx_bytes):
        """Test DOCX extraction giữ cấu trúc text."""
        result = extractor.extract(sample_docx_bytes, "docx")
        # Should have multiple lines (paragraphs separated)
        lines = [l for l in result.split("\n") if l.strip()]
        assert len(lines) >= 3


# ============================================
# PDF EXTRACTION TESTS
# ============================================

class TestPdfExtraction:
    """Tests cho PDF file extraction."""

    def test_extract_pdf(self, extractor, sample_pdf_bytes):
        """Test extract PDF file."""
        result = extractor.extract(sample_pdf_bytes, "pdf")
        assert "Machine Learning" in result
        assert "Supervised Learning" in result

    def test_extract_pdf_content_quality(self, extractor, sample_pdf_bytes):
        """Test PDF extraction chất lượng text."""
        result = extractor.extract(sample_pdf_bytes, "pdf")
        # Text should be reasonably clean
        assert "\x00" not in result  # No null bytes


# ============================================
# VALIDATION TESTS
# ============================================

class TestTextValidation:
    """Tests cho text validation logic."""

    def test_validate_sufficient_text(self, extractor):
        """Test validation với text đủ dài."""
        text = "A" * 200
        is_valid, msg = extractor.validate_extracted_text(text)
        assert is_valid is True
        assert msg == ""

    def test_validate_short_text(self, extractor):
        """Test validation với text quá ngắn."""
        text = "Short text"
        is_valid, msg = extractor.validate_extracted_text(text)
        assert is_valid is False
        assert "too short" in msg

    def test_validate_empty_text(self, extractor):
        """Test validation với text rỗng."""
        is_valid, msg = extractor.validate_extracted_text("")
        assert is_valid is False
        assert "empty" in msg.lower() or "could not" in msg.lower()

    def test_validate_none_text(self, extractor):
        """Test validation với None."""
        is_valid, msg = extractor.validate_extracted_text(None)
        assert is_valid is False


# ============================================
# EDGE CASES & ERROR HANDLING
# ============================================

class TestEdgeCases:
    """Tests cho edge cases và error handling."""

    def test_unsupported_file_type(self, extractor):
        """Test extract file type không hỗ trợ."""
        with pytest.raises(ValueError, match="Unsupported file type"):
            extractor.extract(b"content", "xlsx")

    def test_file_type_case_insensitive(self, extractor, sample_txt_bytes):
        """Test file type không phân biệt hoa thường."""
        result = extractor.extract(sample_txt_bytes, "TXT")
        assert len(result) > 0

    def test_file_type_with_whitespace(self, extractor, sample_txt_bytes):
        """Test file type có khoảng trắng."""
        result = extractor.extract(sample_txt_bytes, "  txt  ")
        assert len(result) > 0

    def test_corrupt_docx(self, extractor):
        """Test extract DOCX file corrupt."""
        with pytest.raises(ValueError, match="Failed to extract"):
            extractor.extract(b"not a valid docx", "docx")

    def test_clean_text_removes_null(self, extractor):
        """Test cleaning removes null characters."""
        text = "Hello\x00World\x00Test content that is long enough to be valid"
        cleaned = extractor._clean_text(text)
        assert "\x00" not in cleaned
        assert "HelloWorldTest" in cleaned.replace(" ", "")

    def test_clean_text_normalizes_newlines(self, extractor):
        """Test cleaning normalizes different newline styles."""
        text = "Line 1\r\nLine 2\rLine 3\nLine 4"
        cleaned = extractor._clean_text(text)
        assert "\r" not in cleaned
        assert "Line 1" in cleaned
        assert "Line 4" in cleaned
