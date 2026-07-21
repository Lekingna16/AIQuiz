"""
Text Extraction Service - Trích xuất văn bản từ file upload
=============================================================

Hỗ trợ 3 định dạng:
- PDF: dùng PyMuPDF (fitz) — nhanh, chính xác, hỗ trợ Unicode/tiếng Việt tốt
- DOCX: dùng python-docx — đọc paragraphs + tables
- TXT: dùng chardet để tự detect encoding

Flow:
  file_bytes + file_type → extract() → cleaned plain text

Edge Cases xử lý:
- PDF scan (chỉ chứa hình ảnh) → trả text rỗng → caller handle
- File DOCX corrupt → raise ValueError
- File TXT encoding lạ → chardet auto-detect
"""

import re
from io import BytesIO

import chardet
import fitz  # PyMuPDF
from docx import Document

# TCVN3 to Unicode converter mapping
TCVN3_CHARS = "µ¸¶·¹¨»¾¼½Æ©ÇÊÈÉË®ÌÐÎÏÑªÒÕÓÔÖ×ÝØÜÞßãáâä«åèæçé¬êíëìîïóñòô\u00adõøö÷ùúýûüþ¡¢§£¤¥¦"
UNICODE_CHARS = "àáảãạăằắẳẵặâầấẩẫậđèéẻẽẹêềếểễệìíỉĩịòóỏõọôồốổỗộơờớởỡợùúủũụưừứửữựỳýỷỹỵĂÂĐÊÔƠƯ"
TCVN3_EXTRA = "\u03bc\u2212"
UNICODE_EXTRA = "àu"

tcvn3_to_unicode_map = dict(zip(TCVN3_CHARS + TCVN3_EXTRA, UNICODE_CHARS + UNICODE_EXTRA))

def _is_tcvn3(text: str) -> bool:
    # Check for characteristic TCVN3 sequences
    indicators = [
        r"[Cc]©u",               # C©u / c©u -> Câu / câu
        r"®\u2212",              # ®− -> đư
        r"l\u03bc",              # lμ -> là
        r"t\u2212\u00ebng",      # t−ëng -> tưởng
        r"tr\u2212\u00eang",     # tr−êng -> trường
        r"th\u2212\u00ebng",     # th−ëng -> thưởng
        r"®êi",                  # ®êi -> đời
        r"®©u",                  # ®©u -> đâu
    ]
    pattern = re.compile("|".join(indicators))
    return bool(pattern.search(text))

def _tcvn3_to_unicode(text: str) -> str:
    pattern = re.compile("|".join(map(re.escape, tcvn3_to_unicode_map.keys())))
    return pattern.sub(lambda m: tcvn3_to_unicode_map[m.group(0)], text)


class TextExtractor:
    """
    Service trích xuất text từ file bytes.

    Usage:
        extractor = TextExtractor()
        text = extractor.extract(file_bytes, "pdf")
    """

    # Mapping file type → extractor method
    _EXTRACTORS = {
        "pdf": "_extract_pdf",
        "docx": "_extract_docx",
        "txt": "_extract_txt",
    }

    # Minimum text length để coi là có nội dung hợp lệ (100 ký tự ≈ 1-2 câu)
    MIN_TEXT_LENGTH = 100

    def extract(self, file_bytes: bytes, file_type: str) -> str:
        """
        Trích xuất text từ file bytes.

        Args:
            file_bytes: Nội dung file dưới dạng bytes
            file_type: Loại file ("pdf", "docx", "txt")

        Returns:
            Cleaned plain text string

        Raises:
            ValueError: Nếu file_type không được hỗ trợ
            ValueError: Nếu không extract được text (file rỗng/scan PDF)
        """
        file_type = file_type.lower().strip()

        extractor_name = self._EXTRACTORS.get(file_type)
        if not extractor_name:
            supported = ", ".join(self._EXTRACTORS.keys())
            raise ValueError(
                f"Unsupported file type: '{file_type}'. "
                f"Supported types: {supported}"
            )

        # Gọi method tương ứng
        extractor = getattr(self, extractor_name)
        raw_text = extractor(file_bytes)

        # Clean text: bỏ whitespace thừa, normalize line breaks
        cleaned = self._clean_text(raw_text)

        return cleaned

    def validate_extracted_text(self, text: str) -> tuple[bool, str]:
        """
        Validate text đã extract có đủ nội dung để tạo quiz hay không.

        Returns:
            (is_valid, error_message)
        """
        if not text or not text.strip():
            return False, "Could not extract any text from the file. The file may be empty or contain only images."

        if len(text.strip()) < self.MIN_TEXT_LENGTH:
            return False, (
                f"Extracted text is too short ({len(text.strip())} characters). "
                f"Minimum {self.MIN_TEXT_LENGTH} characters required to generate meaningful questions."
            )

        return True, ""

    # ============================================
    # Private extraction methods
    # ============================================

    @staticmethod
    def _extract_pdf(file_bytes: bytes) -> str:
        """
        Trích xuất text từ PDF dùng PyMuPDF.

        PyMuPDF (fitz) ưu điểm:
        - Nhanh hơn pdfminer, pdfplumber ~3-5x
        - Xử lý Unicode/tiếng Việt tốt
        - Giữ layout text hợp lý
        - Ít dependency

        Lưu ý: KHÔNG hỗ trợ OCR. Nếu PDF là scan (image-based),
        sẽ trả về text rỗng → caller cần handle case này.
        """
        try:
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            pages_text = []

            for page_num, page in enumerate(doc):
                page_text = page.get_text("text")  # "text" mode: plain text
                if page_text.strip():
                    pages_text.append(page_text)

            doc.close()
            return "\n\n".join(pages_text)

        except Exception as e:
            raise ValueError(f"Failed to extract text from PDF: {str(e)}")

    @staticmethod
    def _extract_docx(file_bytes: bytes) -> str:
        """
        Trích xuất text từ DOCX dùng python-docx.

        Đọc cả paragraphs và tables để không bỏ sót nội dung.
        Nhiều tài liệu giáo dục dùng tables cho bảng so sánh,
        danh sách → cần extract cả phần này.
        """
        try:
            doc = Document(BytesIO(file_bytes))
            parts = []

            # Extract paragraphs
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    parts.append(text)

            # Extract tables (nếu có)
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip()
                        for cell in row.cells
                        if cell.text.strip()
                    )
                    if row_text:
                        parts.append(row_text)

            return "\n".join(parts)

        except Exception as e:
            raise ValueError(f"Failed to extract text from DOCX: {str(e)}")

    @staticmethod
    def _extract_txt(file_bytes: bytes) -> str:
        """
        Đọc file TXT với auto-detect encoding.

        Dùng chardet để xác định encoding thay vì assume UTF-8.
        Nhiều file TXT từ Windows dùng cp1252, GBK, etc.
        """
        try:
            # Detect encoding
            detected = chardet.detect(file_bytes)
            encoding = detected.get("encoding", "utf-8") or "utf-8"
            confidence = detected.get("confidence", 0)

            # Nếu confidence thấp, fallback về utf-8
            if confidence < 0.5:
                encoding = "utf-8"

            return file_bytes.decode(encoding, errors="replace")

        except Exception as e:
            raise ValueError(f"Failed to read TXT file: {str(e)}")

    # ============================================
    # Text cleaning
    # ============================================

    @staticmethod
    def _clean_text(text: str) -> str:
        """
        Clean và normalize text sau khi extract.

        - Bỏ multiple blank lines liên tiếp (giữ tối đa 2)
        - Trim whitespace đầu/cuối mỗi line
        - Bỏ null characters và control characters
        """
        if not text:
            return ""

        # Bỏ null characters (thường xuất hiện trong PDF corrupt)
        text = text.replace("\x00", "")

        # Normalize line endings
        text = text.replace("\r\n", "\n").replace("\r", "\n")

        # Trim mỗi line và bỏ blank lines thừa
        lines = text.split("\n")
        cleaned_lines = []
        blank_count = 0

        for line in lines:
            stripped = line.strip()
            if not stripped:
                blank_count += 1
                if blank_count <= 2:  # Giữ tối đa 2 blank lines liên tiếp
                    cleaned_lines.append("")
            else:
                blank_count = 0
                cleaned_lines.append(stripped)

        return "\n".join(cleaned_lines).strip()
